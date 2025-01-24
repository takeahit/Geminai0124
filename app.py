import re
from io import BytesIO
from typing import List, Tuple, Optional

import pandas as pd
import streamlit as st
import streamlit.components.v1 as components
from docx import Document
from rapidfuzz import fuzz, process
from pydocx import PyDocX
from PyPDF2 import PdfReader
import unicodedata

# --- 定数 ---
MAX_FILE_SIZE = 200 * 1024 * 1024  # 200MB
HIGHLIGHT_COLOR = 6


# --- エラーハンドリングとログ記録 ---
def log_error(message: str):
    """エラーメッセージをログに記録し、Streamlitに表示します."""
    st.error(message)
    # 必要であればログファイルにも出力する

# --- データクリーニング ---
def clean_strings(df: pd.DataFrame) -> pd.DataFrame:
    """文字列データから制御文字を削除します."""
    def clean_cell(value: str) -> str:
        if isinstance(value, str):
            return re.sub(r"[\x00-\x1F\x7F]", "", value)
        return value

    def remove_control_characters(text):
        if isinstance(text, str):
            return ''.join(ch for ch in text if unicodedata.category(ch)[0] != 'C')
        return text

    df = df.applymap(clean_cell)
    df = df.applymap(remove_control_characters)
    return df


def find_invalid_chars(df: pd.DataFrame) -> List[Tuple[str, int, str]]:
    """データフレーム内の非互換文字を検出します."""
    invalid_rows = []
    for col in df.columns:
        for idx, value in df[col].items():
            if isinstance(value, str) and re.search(r"[\x00-\x1F\x7F]", value):
                invalid_rows.append((col, idx, value))
    return invalid_rows

# --- ファイル読み込み ---
def load_excel(file) -> Optional[pd.DataFrame]:
    """Excelファイルを読み込みます。エラー時はNoneを返します."""
    try:
        df = pd.read_excel(file, engine="openpyxl")
        if df.columns.size < 1:
            raise ValueError("Excelファイルには少なくとも1列以上の用語が必要です。")
        return df
    except Exception as e:
        log_error(f"Excelファイルの読み込み中にエラーが発生しました: {e}")
        return None

def extract_text_from_file(file, file_type: str) -> str:
    """ファイルからテキストを抽出します。エラー時は空文字列を返します."""
    try:
        if file_type == "docx":
            doc = Document(file)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        elif file_type == "doc":
             text = PyDocX.to_text(file)
        elif file_type == "pdf":
            reader = PdfReader(file)
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                page_text = page_text.replace("\n", " ").replace("\r", " ")
                page_text = " ".join(page_text.split())
                text += page_text + " "
            text = text.strip()
        else:
             return ""

        # ここでテキストをクリーンアップ
        text = ''.join(ch for ch in text if unicodedata.category(ch)[0] != 'C')
        return text

    except Exception as e:
        log_error(f"ファイルからのテキスト抽出中にエラーが発生しました: {e}")
        return ""

# --- Fuzzy Matching ---
def find_similar_terms(
    text: str, terms: List[str], threshold: int
) -> List[Tuple[str, str]]:
    """テキスト内で類似する用語を検出します."""
    words = text.split()
    detected_terms = []
    for word in words:
        matches = process.extract(word, terms, scorer=fuzz.partial_ratio, limit=10)
        for match in matches:
            if threshold <= match[1] < 100:
                 detected_terms.append((word, match[0]))
    return detected_terms


# --- 修正処理 ---
def apply_corrections_with_table(text, correction_df) -> Tuple[str, List[Tuple[str, str]], int]:
    """正誤表を適用し、修正箇所と回数を記録します."""
    corrections = []
    total_replacements = 0
    for _, row in correction_df.iterrows():
        incorrect, correct = row.iloc[0], row.iloc[1]
        start_index = 0
        while True:
            index = text.find(incorrect, start_index)
            if index == -1:
                break
            text = text[:index] + correct + text[index + len(incorrect):]
            corrections.append((incorrect, correct))
            total_replacements += 1
            start_index = index + len(correct)
    return text, corrections, total_replacements

def apply_kanji_table(text, kanji_df) -> Tuple[str, List[Tuple[str, str]], int]:
    """利用漢字表を適用し、修正箇所と回数を記録します."""
    corrections = []
    total_replacements = 0
    for _, row in kanji_df.iterrows():
        hiragana, kanji = row.iloc[0], row.iloc[1]
        start_index = 0
        while True:
            index = text.find(hiragana, start_index)
            if index == -1:
                break
            text = text[:index] + kanji + text[index + len(hiragana):]
            corrections.append((hiragana, kanji))
            total_replacements += 1
            start_index = index + len(kanji)
    return text, corrections, total_replacements


def create_corrected_word_file_with_formatting(
    original_text: str, corrections: List[Tuple[str, str]]
) -> BytesIO:
    """修正を適用したWordファイルを生成します."""
    doc = Document()
    for paragraph_text in original_text.split("\n"):
        paragraph = doc.add_paragraph()
        start_index = 0
        for incorrect, correct in corrections:
            while incorrect in paragraph_text[start_index:]:
                start_index = paragraph_text.find(incorrect, start_index)
                end_index = start_index + len(incorrect)
                paragraph.add_run(paragraph_text[:start_index])
                run = paragraph.add_run(correct)
                run.font.highlight_color = HIGHLIGHT_COLOR
                paragraph_text = paragraph_text[end_index:]
                start_index = 0
        try:
            paragraph.add_run(paragraph_text)
        except UnicodeEncodeError as e:
            log_error(f"テキストの追加中にエンコードエラーが発生しました: {e}. 残りのテキストをスキップします。")
            paragraph.add_run(paragraph_text.encode('unicode_escape').decode('utf-8')) # エンコードエラー発生時、エスケープ処理を行う
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# --- データ表示とダウンロード ---
def create_correction_table(detected: List[Tuple[str, str]]) -> pd.DataFrame:
    """検出された類似語をデータフレームに変換します."""
    if not detected:
        return pd.DataFrame(columns=["原稿内の語", "類似する用語"])
    return pd.DataFrame(detected, columns=["原稿内の語", "類似する用語"])

def download_excel(df: pd.DataFrame, file_name: str, sheet_name: str):
    """データフレームをExcelファイルとしてダウンロードします."""
    output = BytesIO()
    with pd.ExcelWriter(output) as writer:
        df.to_excel(writer, index=False, sheet_name=sheet_name)
    st.download_button(
        label=f"{sheet_name}をダウンロード",
        data=output.getvalue(),
        file_name=file_name,
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

def download_word(file: BytesIO, file_name: str):
    """Wordファイルをダウンロードします."""
    st.download_button(
        label="修正済みファイルをダウンロード",
        data=file.getvalue(),
        file_name=file_name,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

# --- メイン処理関数 ---
# Streamlit アプリケーション
st.markdown("<h1 style='text-align: center;'>南江堂用用語チェッカー（笑）</h1>", unsafe_allow_html=True)

# 左右のカラムを作成
col1, col2 = st.columns([1, 1])

# 左側のカラム（Difyチャットボット）
with col1:
    iframe_html = """
        <iframe
          src="https://udify.app/chatbot/rGMuWhHEu9Hcwbqe"
          style="width: 100%; height: 700px; min-height: 700px"
          frameborder="0"
          allow="microphone">
        </iframe>
    """
    components.html(iframe_html, height=700)

# 右側のカラム（ファイルアップローダー）
with col2:
    st.write("以下のファイルを個別にアップロードしてください:")
    word_file = st.file_uploader("原稿ファイル (Word, DOC, PDF):", type=["docx", "doc", "pdf"])
    terms_file = st.file_uploader("用語集ファイル (A列に正しい用語を記載したExcel):", type=["xlsx"])
    correction_file = st.file_uploader("正誤表ファイル (A列に誤った用語、B列に正しい用語を記載したExcel):", type=["xlsx"])
    kanji_file = st.file_uploader("利用漢字表ファイル (A列にひらがな、B列に漢字を記載したExcel):", type=["xlsx"])

    # アップロードファイルサイズの制限 (200MB以下)
    for file, name in [
        (word_file, "原稿ファイル"),
        (terms_file, "用語集ファイル"),
        (correction_file, "正誤表ファイル"),
        (kanji_file, "利用漢字表ファイル"),
    ]:
        if file and file.size > MAX_FILE_SIZE:
            st.error(f"{name}のサイズが大きすぎます（{MAX_FILE_SIZE / (1024 * 1024)}MB以下にしてください）。")
            st.stop()

    if word_file and (terms_file or correction_file or kanji_file):
        file_type = word_file.name.split(".")[-1]
        original_text = extract_text_from_file(word_file, file_type)

        corrections = []
        total_replacements = 0

        if terms_file:
             try:
                terms_df = load_excel(terms_file)
                if terms_df is not None:
                    terms_df = clean_strings(terms_df)
                    terms = terms_df.iloc[:, 0].dropna().astype(str).tolist()
                    threshold = st.slider("類似度の閾値を設定してください (50-99):", min_value=50, max_value=99, value=65)
                    detected = find_similar_terms(original_text, terms, threshold)

                    if detected:
                        st.success(f"類似語が{len(detected)}件検出されました！")
                        correction_table = create_correction_table(detected)
                        st.dataframe(correction_table)

                        output = BytesIO()
                        with pd.ExcelWriter(output) as writer:
                            correction_table.to_excel(writer, index=False, sheet_name="修正箇所一覧")
                        st.download_button(
                            label="修正箇所一覧をダウンロード",
                            data=output.getvalue(),
                            file_name="修正箇所一覧.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        )

             except Exception as e:
                   st.error(f"用語集ファイルの処理中にエラーが発生しました: {e}")
        if correction_file:
            try:
                correction_df = load_excel(correction_file)
                if correction_df is not None:
                    correction_df = clean_strings(correction_df)
                    updated_text, corrections_from_table, replacements = apply_corrections_with_table(original_text, correction_df)
                    corrections.extend(corrections_from_table)
                    total_replacements += replacements

                    st.success(f"正誤表を適用し、{replacements}回の修正を行いました！")

                    corrections_df = pd.DataFrame(corrections_from_table, columns=["誤った用語", "正しい用語"])
                    st.dataframe(corrections_df)

                    output = BytesIO()
                    with pd.ExcelWriter(output) as writer:
                        corrections_df.to_excel(writer, index=False, sheet_name="正誤表修正箇所")
                    st.download_button(
                        label="正誤表修正箇所をダウンロード",
                        data=output.getvalue(),
                        file_name="正誤表修正箇所.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    )

                    corrected_file = create_corrected_word_file_with_formatting(original_text, corrections_from_table)
                    st.download_button(
                        label="正誤表修正済みファイルをダウンロード",
                        data=corrected_file.getvalue(),
                        file_name="正誤表修正済み.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )

            except Exception as e:
                    st.error(f"正誤表ファイルの処理中にエラーが発生しました: {e}")

        if kanji_file:
                try:
                    kanji_df = load_excel(kanji_file)
                    if kanji_df is not None:
                        kanji_df = clean_strings(kanji_df)
                        updated_text, kanji_corrections, replacements = apply_kanji_table(original_text, kanji_df)
                        corrections.extend(kanji_corrections)
                        total_replacements += replacements

                        st.success(f"利用漢字表を適用し、{replacements}回の修正を行いました！")

                        kanji_corrections_df = pd.DataFrame(kanji_corrections, columns=["ひらがな", "漢字"])
                        st.dataframe(kanji_corrections_df)

                        output = BytesIO()
                        with pd.ExcelWriter(output) as writer:
                            kanji_corrections_df.to_excel(writer, index=False, sheet_name="漢字修正箇所")
                        st.download_button(
                            label="利用漢字表修正箇所をダウンロード",
                            data=output.getvalue(),
                            file_name="利用漢字表修正箇所.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        )

                        corrected_file = create_corrected_word_file_with_formatting(original_text, kanji_corrections)
                        st.download_button(
                            label="利用漢字表修正済みファイルをダウンロード",
                            data=corrected_file.getvalue(),
                            file_name="利用漢字表修正済み.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )

                except Exception as e:
                    st.error(f"利用漢字表ファイルの処理中にエラーが発生しました: {e}")

        st.markdown(f"<h3 style='text-align: left;'>正誤表と利用漢字表を適用し、{total_replacements}回の修正を行いました！</h3>", unsafe_allow_html=True)

    else:
        st.warning("原稿ファイルと、用語集、正誤表、利用漢字表のいずれかをアップロードしてください！")
